# processor.py

import re
from typing import List, Dict
from io import BytesIO

# Document processing imports
import pymupdf  # PyMuPDF (fitz)
from docx import Document as DocxDocument
from fastapi import UploadFile

# Text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter


class EnhancedDocumentProcessor:
    """
    Enhanced document processor focused on loading and chunking.
    """

    def __init__(self):
        pass

    # ============= ENHANCED DOCUMENT LOADING =============

    async def load_document(self, file: UploadFile, enable_clause_grouping: bool = True) -> str:
        """Enhanced document loading with support for multiple formats."""
        contents = await file.read()
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            text = self.extract_text_from_pdf(contents)
        elif filename.endswith(".docx"):
            text = self.extract_text_from_docx(contents)
        else:
            text = contents.decode()

        # Apply clause grouping if enabled
        if enable_clause_grouping:
            text = self.process_text_with_clause_grouping(text)

        return text

    def is_footer_content(self, text: str, page_height: float = None, y_position: float = None) -> bool:
        """Detect if text is likely footer content."""
        if not text.strip():
            return True
        text_lower = text.lower().strip()
        footer_patterns = [
            r'page\s+\d+', r'\d+\s*/\s*\d+', r'©\s*\d{4}', r'copyright\s+\d{4}',
            r'confidential', r'proprietary', r'all rights reserved', r'footer',
            r'^\d+$', r'^\s*-\s*\d+\s*-\s*$'
        ]
        if any(re.search(pattern, text_lower) for pattern in footer_patterns):
            return True
        if page_height and y_position and (y_position > (page_height * 0.9)):
            return True
        if len(text.strip()) < 10 and (text_lower.isdigit() or any(word in text_lower for word in ['page', 'copyright', '©'])):
            return True
        return False

    def extract_text_from_pdf(self, data: bytes) -> str:
        """Enhanced PDF extraction with footer filtering and structure preservation."""
        doc = pymupdf.open(stream=data, filetype="pdf")
        extracted = []
        for page in doc:
            page_height = page.rect.height
            sorted_blocks = sorted(page.get_text("blocks"), key=lambda b: (b[1], b[0]))
            page_content = []
            for block in sorted_blocks:
                _, _, _, y1, block_text, _, _ = block
                if self.is_footer_content(block_text, page_height, y1):
                    continue
                block_text = block_text.strip()
                if block_text:
                    block_text = re.sub(r'\s+', ' ', block_text)
                    block_text = re.sub(r'-\s*\n\s*', '', block_text)
                    page_content.append(block_text)
            
            if page_content:
                page_text = "\n\n".join(page_content)
                formatted_page = self.format_structured_content(page_text)
                if formatted_page.strip():
                    extracted.append(formatted_page)
        doc.close()
        return "\n\n---PAGE_BREAK---\n\n".join(extracted)

    def extract_text_from_docx(self, data: bytes) -> str:
        """Enhanced DOCX extraction."""
        file_stream = BytesIO(data)
        doc = DocxDocument(file_stream)
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and not self.is_footer_content(text):
                text_parts.append(self.format_paragraph_structure(text, para))
        for table in doc.tables:
            table_text = self.extract_table_properly(table, flatten=True)
            if table_text:
                text_parts.append(table_text)
        final_text = "\n\n".join(text_parts)
        return self.format_structured_content(final_text)

    # ============= STRUCTURE DETECTION & FORMATTING =============

    def format_paragraph_structure(self, text: str, paragraph=None) -> str:
        """Format paragraph with structure detection."""
        if re.match(r'^\s*([0-9]+[.\)]|[a-z][.\)]|[ivx]+[.\)])\s+', text, re.IGNORECASE):
            return f"LIST_ITEM: {text}"
        if re.match(r'^\s*[•·▪▫◦‣⁃-]\s+', text):
            return f"BULLET_ITEM: {text}"
        if len(text) < 100 and not text.endswith(('.', ',')):
            if text.isupper() or (len(text.split()) <= 10):
                return f"HEADING: {text}"
        return text

    def extract_table_properly(self, table, flatten: bool = True) -> str:
        """Extract table with proper formatting and optional flattening."""
        if not table.rows:
            return ""
        table_data = []
        for row in table.rows:
            row_cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            if any(cell.strip() for cell in row_cells):
                table_data.append(row_cells)
        if not table_data:
            return ""
        if flatten:
            flattened_sentences = self.flatten_table(table_data)
            if flattened_sentences:
                return f"TABLE_START\n" + "\n".join(flattened_sentences) + "\nTABLE_END"
        else:
            table_lines = [" || ".join(row) for row in table_data]
            return f"TABLE_START\n" + "\n".join(table_lines) + "\nTABLE_END"
        return ""

    def format_structured_content(self, text: str) -> str:
        """Format structured content like lists, tables, and sections."""
        lines = text.split('\n')
        formatted_lines = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if self.is_structured_section_start(line):
                section = self.extract_structured_section(lines, i)
                formatted_lines.append("STRUCTURED_SECTION_START")
                formatted_lines.append(f"SECTION_HEADER: {lines[i].strip()}")
                formatted_lines.extend(section['items'])
                formatted_lines.append("STRUCTURED_SECTION_END")
                i = section['next_index']
            elif self.is_tabular_line(line):
                table_section = self.extract_tabular_section(lines, i)
                if len(table_section['lines']) > 1:
                    formatted_lines.append("TABLE_START")
                    for table_line in table_section['lines']:
                        formatted_lines.append(re.sub(r'\s{3,}', ' || ', table_line))
                    formatted_lines.append("TABLE_END")
                    i = table_section['next_index']
                else:
                    formatted_lines.append(line)
                    i += 1
            else:
                formatted_lines.append(line)
                i += 1
        return '\n'.join(formatted_lines)

    def is_structured_section_start(self, line: str) -> bool:
        """Detect if a line starts a structured section."""
        line_lower = line.lower()
        section_indicators = [
            'waiting period', 'conditions', 'benefits', 'coverage', 'exclusions',
            'features', 'limits', 'terms', 'definitions', 'procedures'
        ]
        if any(indicator in line_lower for indicator in section_indicators):
            return True
        if re.match(r'^\s*([0-9]+\.|\([0-9]+\)|[a-z]\.|\([a-z]\))', line_lower):
            return True
        return False

    def extract_structured_section(self, lines: List[str], start_index: int) -> Dict:
        """Extract a structured section with its items."""
        items = []
        current_index = start_index + 1
        while current_index < len(lines):
            line = lines[current_index].strip()
            if not line:
                current_index += 1
                continue
            if re.match(r'^\s*([a-z]\.|\d+\.|[ivx]+\.)', line.lower()):
                items.append(f"ITEM: {line}")
                current_index += 1
            elif line.lower().startswith(('note:', 'above', 'these', 'all')):
                items.append(f"NOTE: {line}")
                current_index += 1
                break
            else:
                break
        return {'items': items, 'next_index': current_index}

    def is_tabular_line(self, line: str) -> bool:
        """Detect if a line is part of tabular data."""
        if not line.strip(): return False
        currency_pattern = r'(INR|Rs\.?|₹|\$)\s*[\d,]+(\.\d+)?'
        if len(re.findall(currency_pattern, line)) >= 2: return True
        if len(re.findall(r'\d+%', line)) >= 2: return True
        if len(re.findall(r'Up to\s+[\w\s]+', line, re.IGNORECASE)) >= 2: return True
        if len(re.findall(r'Plan\s+[A-Z]', line, re.IGNORECASE)) >= 2: return True
        if len(re.findall(r'\s{3,}', line)) >= 2: return True
        return False

    def extract_tabular_section(self, lines: List[str], start_index: int) -> Dict:
        """Extract a tabular section."""
        table_lines = []
        current_index = start_index
        while current_index < len(lines):
            line = lines[current_index].strip()
            if self.is_tabular_line(line):
                table_lines.append(line)
                current_index += 1
            elif not line:
                if len([l for l in lines[current_index:current_index+3] if l.strip()]) == 0:
                    break
                current_index += 1
            else:
                break
        return {'lines': table_lines, 'next_index': current_index}

    # ============= ENHANCED CHUNKING STRATEGIES =============

    def chunk_page_content(self, content: str, max_size: int, overlap: int) -> List[str]:
        """Chunk page content while preserving structure."""
        if len(content) <= max_size:
            return [content]
        chunks = []
        sections = re.split(r'((?:STRUCTURED_SECTION_START.*?STRUCTURED_SECTION_END|TABLE_START.*?TABLE_END))', content, flags=re.DOTALL)
        current_chunk = ""
        for section in sections:
            if not section.strip():
                continue
            is_structured = section.startswith(('STRUCTURED_SECTION_START', 'TABLE_START'))
            if is_structured:
                if len(current_chunk) + len(section) > max_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = section
                else:
                    current_chunk += ("\n\n" if current_chunk else "") + section
            else:
                paragraphs = [p for p in section.split('\n\n') if p.strip()]
                for para in paragraphs:
                    if len(current_chunk) + len(para) > max_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
                            current_chunk = overlap_text + "\n\n" + para
                        else:
                            current_chunk = para
                    else:
                        current_chunk += ("\n\n" if current_chunk else "") + para
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        return chunks
        
    def flatten_table(self, table_data: List[List[str]]) -> List[str]:
        """Flatten table to natural language sentences."""
        if not table_data or len(table_data) < 2:
            return []
        headers = [str(h).strip() for h in table_data[0]]
        flattened_rows = []
        for row in table_data[1:]:
            subject = str(row[0]).strip()
            if not subject:
                continue
            for i, cell in enumerate(row[1:]):
                cell_text = str(cell).strip()
                header = headers[i + 1] if i + 1 < len(headers) else f"Column_{i+1}"
                if cell_text and header:
                    sentence = f"For the benefit '{subject}', the value under '{header}' is '{cell_text}'."
                    flattened_rows.append(sentence)
        return flattened_rows

    def detect_content_type(self, text: str) -> str:
        """Detect the primary content type of a text chunk."""
        text_lower = text.lower().strip()
        if 'table_start' in text_lower or 'for the benefit' in text_lower:
            return 'table'
        clause_patterns = [
            r'clause_start:', r'sub_clause:', r'clause\s+\d+', r'section\s+\d+', r'article\s+\d+',
            r'paragraph\s+\d+', r'subsection', r'definition[s]?:', r'terms\s+and\s+conditions',
            r'coverage\s+details', r'exclusions?:', r'waiting\s+period', r'deductible', r'premium'
        ]
        if any(re.search(pattern, text_lower) for pattern in clause_patterns):
            return 'clause'
        if any(marker in text_lower for marker in ['list_item:', 'bullet_item:', 'structured_section_start']):
            return 'structured_list'
        return 'text'

    def process_text_with_clause_grouping(self, text: str) -> str:
        """Process text with optional clause grouping."""
        pages = text.split('---PAGE_BREAK---')
        processed_pages = []
        for page in pages:
            if not page.strip(): continue
            clauses = self.group_clauses(page)
            if clauses and len(clauses) > 1:
                grouped_text = '\n\n'.join(clauses)
                processed_page = self.add_clause_markers(grouped_text)
                processed_pages.append(processed_page)
            else:
                processed_pages.append(page)
        return '\n\n---PAGE_BREAK---\n\n'.join(processed_pages)

    def add_clause_markers(self, text: str) -> str:
        """Add clause markers to grouped clauses for better identification."""
        lines = text.split('\n')
        marked_lines = []
        for line in lines:
            line = line.strip()
            if not line: continue
            if re.match(r'^\d+\)\s.*?\(Code\s-\w+\)', line):
                marked_lines.append(f'CLAUSE_START: {line}')
            elif re.match(r'^[a-z]\.\s', line):
                marked_lines.append(f'SUB_CLAUSE: {line}')
            else:
                marked_lines.append(line)
        return '\n'.join(marked_lines)

    def group_clauses(self, text: str) -> List[str]:
        """Enhanced clause grouping with regex patterns."""
        # NOTE: This is the corrected and single version of this method.
        pattern = r'(\n\d+\)\s.*?\(Code\s-\w+\)|\n[a-z]\.\s)'
        parts = re.split(pattern, text)
        grouped_clauses = []
        if parts[0] and parts[0].strip():
            grouped_clauses.append(parts[0].strip())
        for i in range(1, len(parts), 2):
            heading = parts[i].strip()
            content = parts[i + 1].strip() if i + 1 < len(parts) else ""
            grouped_clauses.append(f"{heading}\n{content}")
        return [clause for clause in grouped_clauses if clause]

    def create_chunk_with_metadata(self, chunk_text: str, chunk_index: int, page_number: int, source_file: str) -> Dict:
        """Create a chunk with comprehensive metadata."""
        clean_text = self.clean_chunk_text(chunk_text)
        return {
            'text': clean_text,
            'raw_text': chunk_text,
            'type': self.detect_content_type(chunk_text),
            'chunk_index': chunk_index,
            'page_number': page_number,
            'source_file': source_file,
            'character_count': len(clean_text),
            'word_count': len(clean_text.split()),
            'has_table': 'table_start' in chunk_text.lower(),
            'has_list': any(marker in chunk_text.lower() for marker in ['list_item:', 'bullet_item:']),
            'has_structured_section': 'structured_section_start' in chunk_text.lower(),
        }

    def clean_chunk_text(self, chunk_text: str) -> str:
        """Clean chunk text by removing structure markers."""
        markers_to_remove = [
            'TABLE_START', 'TABLE_END', 'STRUCTURED_SECTION_START', 'STRUCTURED_SECTION_END',
            'LIST_ITEM: ', 'BULLET_ITEM: ', 'HEADING: ', 'SECTION_HEADER: ',
            'ITEM: ', 'NOTE: ', 'CLAUSE_START: ', 'SUB_CLAUSE: '
        ]
        clean_text = chunk_text
        for marker in markers_to_remove:
            clean_text = clean_text.replace(marker, '')
        clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text.strip())
        return clean_text

    def smart_chunk_text_with_metadata(self, text: str, max_chunk_size: int = 1000, overlap: int = 50, source_file: str = None) -> List[Dict]:
        """Chunk text while preserving structure and adding comprehensive metadata."""
        chunks_with_metadata = []
        pages = text.split('---PAGE_BREAK---')
        chunk_index = 0
        for page_num, page in enumerate(pages, 1):
            if not page.strip():
                continue
            page_chunks = self.chunk_page_content(page.strip(), max_chunk_size, overlap)
            for chunk_text in page_chunks:
                if chunk_text.strip():
                    metadata = self.create_chunk_with_metadata(chunk_text, chunk_index, page_num, source_file)
                    chunks_with_metadata.append(metadata)
                    chunk_index += 1
        return chunks_with_metadata

    # ============= UTILITY METHODS =============

    def get_text_stats(self, text: str) -> Dict:
        """Get statistics about the processed text."""
        return {
            'total_characters': len(text),
            'total_lines': len(text.split('\n')),
            'total_words': len(text.split()),
            'pages': text.count('---PAGE_BREAK---') + 1,
            'tables': text.count('TABLE_START'),
            'structured_sections': text.count('STRUCTURED_SECTION_START'),
        }