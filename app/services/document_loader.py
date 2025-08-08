from fastapi import UploadFile
from docx import Document
from fitz import open as fitz_open
import tempfile
import os
import io
import re
from typing import List, Dict, Tuple
from app.services.parser.excel import extract_text_from_excel_bytes, extract_text_from_image_bytes, extract_text_from_pptx_with_ocr, extract_text_from_csv_bytes, extract_text_from_nested_zip, extract_text_from_txt

from typing import Tuple
import aiohttp
import io

import httpx
import bs4

async def extract_token_from_webpage(url: str) -> str:
    """
    Fetches an HTML webpage and extracts the content of the element with id="token".
    """
    async with httpx.AsyncClient() as client:
        response = await client.get(url)
        response.raise_for_status()

    html = response.text
    soup = BeautifulSoup(html, 'html.parser')
    token_element = soup.find(id="token")

    if token_element:
        return token_element.get_text(strip=True)
    else:
        raise ValueError("Element with id='token' not found in the webpage.")
    


async def download_file(url: str) -> tuple[str, bytes]:
    async with httpx.AsyncClient() as client:
        response = await client.get(url)
        response.raise_for_status()
        filename = url.split("?")[0].split("/")[-1]  # Extract file name
        return filename, response.content




async def load_document(url: str) -> str:
    if "hackrx.in/utils/get-secret-token" in url:
        token = await extract_token_from_webpage(url)
        return token
    
    filename, contents = await download_file(url)
    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(contents)  # contents is bytes
    elif filename.endswith(".docx"):
        return extract_text_from_docx(contents)
    elif filename.endswith(".xls") or filename.endswith(".xlsx"):
        return extract_text_from_excel_bytes(contents)
    elif filename.endswith(".pptx"):
        return extract_text_from_pptx_with_ocr(contents)
    elif filename.endswith(".csv"):
        return extract_text_from_csv_bytes(contents)
    elif filename.endswith(".zip"):
        return extract_text_from_nested_zip(contents)
    elif filename.endswith(".txt"):
        return extract_text_from_txt(contents)
    elif filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".jpeg"):
        return extract_text_from_image_bytes(contents)
    else:
        raise ValueError("Unsupported file format")
    

def is_footer_content(text: str, page_height: float = None, y_position: float = None) -> bool:
    """
    Detect if text is likely footer content based on content patterns and position
    """
    if not text.strip():
        return True
    
    text_lower = text.lower().strip()
    
    # Common footer patterns
    footer_patterns = [
        r'page\s+\d+',  # "Page 1", "Page 2"
        r'\d+\s*/\s*\d+',  # "1/5", "2 / 10"
        r'©\s*\d{4}',  # Copyright year
        r'copyright\s+\d{4}',
        r'confidential',
        r'proprietary',
        r'all rights reserved',
        r'footer',
        r'^\d+$',  # Just page numbers
        r'^\s*-\s*\d+\s*-\s*$',  # "-1-", "- 2 -"
    ]
    
    # Check if text matches footer patterns
    for pattern in footer_patterns:
        if re.search(pattern, text_lower):
            return True
    
    # Check position-based criteria (bottom 10% of page)
    if page_height and y_position:
        if y_position > (page_height * 0.9):
            return True
    
    # Very short lines at document boundaries are often footers
    if len(text.strip()) < 10 and (text_lower.isdigit() or 
                                   any(word in text_lower for word in ['page', 'copyright', '©'])):
        return True
    
    return False


def extract_text_from_pdf(data: bytes) -> str:
    doc = fitz_open(stream=data, filetype="pdf")
    extracted = []

    for page_num, page in enumerate(doc):
        page_height = page.rect.height
        
        # Get text in blocks to preserve layout
        blocks = page.get_text("blocks")
        # Sort blocks by y (top to bottom), then x (left to right)
        sorted_blocks = sorted(blocks, key=lambda b: (b[1], b[0]))
        
        page_content = []
        
        for block in sorted_blocks:
            x0, y0, x1, y1, block_text, block_no, block_type = block
            
            # Skip footer content
            if is_footer_content(block_text, page_height, y1):
                continue
                
            block_text = block_text.strip()
            if not block_text:
                continue
            
            # Clean and normalize the text
            block_text = re.sub(r'\s+', ' ', block_text)  # Normalize whitespace
            block_text = re.sub(r'-\s*\n\s*', '', block_text)  # Remove hyphenation
            
            page_content.append(block_text)
        
        # Format structured content
        page_text = "\n\n".join(page_content)
        formatted_page = format_structured_content(page_text)
        
        if formatted_page.strip():
            extracted.append(formatted_page)

    return "\n\n---PAGE_BREAK---\n\n".join(extracted)


def extract_text_from_docx(data: bytes) -> str:
    file_stream = io.BytesIO(data)
    doc = Document(file_stream)

    text_parts = []

    # Extract paragraphs with structure preservation
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Skip footer content
        # if is_footer_content(text):
        #     continue
        
        # Preserve paragraph structure
        formatted_text = format_paragraph_structure(text, para)
        text_parts.append(formatted_text)

    # Extract tables with proper formatting
    for table in doc.tables:
        table_text = extract_table_properly(table)
        if table_text:
            text_parts.append(table_text)

    final_text = "\n\n".join(text_parts)
    return format_structured_content(final_text)


def format_paragraph_structure(text: str, paragraph) -> str:
    """Format paragraph with structure detection"""
    # Detect different types of content structure
    
    # Numbered/lettered lists (1. 2. or a. b. or i. ii.)
    if re.match(r'^\s*([0-9]+[.\)]|[a-z][.\)]|[ivx]+[.\)])\s+', text, re.IGNORECASE):
        return f"LIST_ITEM: {text}"
    
    # Bullet points
    if re.match(r'^\s*[•·▪▫◦‣⁃-]\s+', text):
        return f"BULLET_ITEM: {text}"
    
    # Headings (basic detection)
    if len(text) < 100 and not text.endswith('.') and not text.endswith(','):
        if text.isupper() or (len(text.split()) <= 10):
            return f"HEADING: {text}"
    
    return text


def extract_table_properly(table) -> str:
    """Extract table with proper formatting"""
    if not table.rows:
        return ""
    
    table_lines = []
    
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for cell in row.cells:
            # Clean cell text and preserve line breaks within cells
            cell_text = cell.text.strip().replace('\n', ' | ')
            row_cells.append(cell_text)
        
        if any(cell.strip() for cell in row_cells):  # Only non-empty rows
            table_lines.append(" || ".join(row_cells))
    
    if table_lines:
        return f"TABLE_START\n" + "\n".join(table_lines) + "\nTABLE_END"
    
    return ""


def format_structured_content(text: str) -> str:
    """Format structured content like lists, tables, and sections"""
    lines = text.split('\n')
    formatted_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            formatted_lines.append("")
            i += 1
            continue
        
        # Detect multi-line structured sections (like waiting periods, benefits, etc.)
        if is_structured_section_start(line):
            section = extract_structured_section(lines, i)
            formatted_lines.append("STRUCTURED_SECTION_START")
            formatted_lines.append(f"SECTION_HEADER: {line}")
            formatted_lines.extend(section['items'])
            formatted_lines.append("STRUCTURED_SECTION_END")
            i = section['next_index']
            continue
        
        # Detect tabular patterns
        if is_tabular_line(line):
            table_section = extract_tabular_section(lines, i)
            if len(table_section['lines']) > 1:  # Only if multiple rows
                formatted_lines.append("TABLE_START")
                for table_line in table_section['lines']:
                    # Format with consistent separators
                    formatted_line = re.sub(r'\s{3,}', ' || ', table_line)
                    formatted_lines.append(formatted_line)
                formatted_lines.append("TABLE_END")
                i = table_section['next_index']
                continue
        
        # Regular line
        formatted_lines.append(line)
        i += 1
    
    return '\n'.join(formatted_lines)


def is_structured_section_start(line: str) -> bool:
    """Detect if line starts a structured section (lists, conditions, etc.)"""
    line_lower = line.lower()
    
    # Common section indicators
    section_indicators = [
        'waiting period', 'conditions', 'benefits', 'coverage', 'exclusions',
        'features', 'limits', 'terms', 'definitions', 'procedures'
    ]
    
    for indicator in section_indicators:
        if indicator in line_lower:
            return True
    
    # Check for numbered/lettered section headers
    if re.match(r'^\s*([0-9]+\.|\([0-9]+\)|[a-z]\.|\([a-z]\))', line_lower):
        return True
    
    return False


def extract_structured_section(lines: List[str], start_index: int) -> Dict:
    """Extract a structured section with its items"""
    items = []
    current_index = start_index + 1
    
    while current_index < len(lines):
        line = lines[current_index].strip()
        
        if not line:
            current_index += 1
            continue
        
        # Check if this is a list item (a. b. c. or 1. 2. 3.)
        if re.match(r'^\s*([a-z]\.|\d+\.)', line.lower()):
            items.append(f"ITEM: {line}")
            current_index += 1
        elif re.match(r'^\s*[ivx]+\.', line.lower()):  # Roman numerals
            items.append(f"ITEM: {line}")
            current_index += 1
        elif line.lower().startswith(('note:', 'above', 'these', 'all')):
            items.append(f"NOTE: {line}")
            current_index += 1
            break
        else:
            # If it doesn't match expected pattern, section is done
            break
    
    return {
        'items': items,
        'next_index': current_index
    }


def is_tabular_line(line: str) -> bool:
    """Detect if line is part of tabular data"""
    if not line.strip():
        return False
    
    # Multiple currency values or amounts
    currency_pattern = r'(INR|Rs\.?|₹|\$)\s*[\d,]+(\.\d+)?'
    if len(re.findall(currency_pattern, line)) >= 2:
        return True
    
    # Multiple percentage values
    if len(re.findall(r'\d+%', line)) >= 2:
        return True
    
    # Pattern like "Up to X Up to Y"
    if len(re.findall(r'Up to\s+[\w\s]+', line, re.IGNORECASE)) >= 2:
        return True
    
    # Multiple plan references (Plan A, Plan B, etc.)
    if len(re.findall(r'Plan\s+[A-Z]', line, re.IGNORECASE)) >= 2:
        return True
    
    # Consistent spacing suggesting columns (3+ spaces between words)
    if len(re.findall(r'\s{3,}', line)) >= 2:
        return True
    
    return False


def extract_tabular_section(lines: List[str], start_index: int) -> Dict:
    """Extract a tabular section"""
    table_lines = []
    current_index = start_index
    
    while current_index < len(lines):
        line = lines[current_index].strip()
        
        if not line:
            current_index += 1
            if len([l for l in lines[current_index:current_index+3] if l.strip()]) == 0:
                break  # Multiple empty lines = end of table
            continue
        
        if is_tabular_line(line):
            table_lines.append(line)
            current_index += 1
        else:
            break
    
    return {
        'lines': table_lines,
        'next_index': current_index
    }


def smart_chunk_text(text: str, max_chunk_size: int = 1000, overlap: int = 50) -> List[str]:
    """
    Chunk text while preserving structure
    """
    chunks = []
    
    # Split by page breaks first
    pages = text.split('---PAGE_BREAK---')
    
    for page in pages:
        if not page.strip():
            continue
        
        page_chunks = chunk_page_content(page.strip(), max_chunk_size, overlap)
        chunks.extend(page_chunks)
    
    return chunks


def chunk_page_content(content: str, max_size: int, overlap: int) -> List[str]:
    """Chunk page content while preserving structure"""
    if len(content) <= max_size:
        return [content]
    
    chunks = []
    
    # Split by structured sections and tables
    sections = re.split(r'((?:STRUCTURED_SECTION_START.*?STRUCTURED_SECTION_END|TABLE_START.*?TABLE_END))', 
                       content, flags=re.DOTALL)
    
    current_chunk = ""
    
    for section in sections:
        if not section.strip():
            continue
        
        # Keep structured sections and tables together
        if (section.startswith('STRUCTURED_SECTION_START') or 
            section.startswith('TABLE_START')):
            
            if len(current_chunk) + len(section) > max_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = section
            else:
                current_chunk += "\n\n" + section if current_chunk else section
        else:
            # Handle regular text
            paragraphs = [p for p in section.split('\n\n') if p.strip()]
            
            for para in paragraphs:
                if len(current_chunk) + len(para) > max_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        # Add minimal overlap
                        overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
                        current_chunk = overlap_text + "\n\n" + para
                    else:
                        current_chunk = para
                else:
                    current_chunk += "\n\n" + para if current_chunk else para
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks